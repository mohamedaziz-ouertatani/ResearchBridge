"""Render a ResearchAssessment as a downloadable .docx or .pdf (blueprint Sec 2A).

Both formats print the same content in the same order as the web report
(AssessmentReport.tsx) and by the same design principle: an assessment is
never its own source of truth, so every gradeable field's full supporting
evidence is printed inline here rather than collapsed the way the web UI's
<details> element hides it - the exported file has to stand on its own
without a browser to expand anything in.

build_report_sections() is the one place that decides *what* goes in the
report; build_docx()/build_pdf() only decide *how* to lay it out for their
format. A future third format is a new renderer over the same sections,
not a rework of the content logic.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path

from researchbridge.api.schemas import AnalysisClaimOut, AssessmentEvidenceOut, ResearchAssessmentOut
from researchbridge.assessment.export_charts import (
    INK,
    INK_FAINT,
    INK_SOFT,
    RULE,
    evidence_bar_chart_png,
    level_gauge_png,
)

# The same ink-gray palette and font pairing as the web report (globals.css /
# layout.tsx) - Space Grotesk for headings/labels, Source Serif 4 for body
# text and quotes. No color accents: teal and amber are reserved for cosine
# distance and the active/typing state respectively, neither of which applies
# to a static export. (INK/INK_SOFT/INK_FAINT/RULE live in export_charts so
# the chart helpers share the same palette without importing this module.)

FONTS_DIR = Path(__file__).parent / "fonts"

OPPORTUNITIES_REASON = (
    "Not generated. Naming a product opportunity means inventing a claim the "
    "literature does not make, so this is left to a human reviewer."
)

_GAP_UNASSESSED_REASONS = {
    "no_relevant_evidence": (
        "Not assessed - insufficient relevant evidence was retrieved for this input to "
        "investigate whether a research gap exists."
    ),
    "checked_no_gap_found": "No gap was found in the retrieved literature for this input.",
    None: "No gap was found in the retrieved literature for this input.",
}

_COMPARISON_CLAIM_RE = re.compile(r'^-\s*"([^"]*)":\s*(.*)$')
"""Matches one line of comparison_summary, same pattern as the web report's
ComparisonSummary component (AssessmentReport.tsx)."""

_APPLICATIONS_UNASSESSED_REASONS = {
    "not_assessed": "No relevant paper was retrieved for this input, so applications could not be assessed.",
    "no_evidence": "No retrieved paper stated an application.",
}


@dataclass
class ReportSection:
    label: str
    level: str | None = None
    """A rule-based category (novelty/feasibility level), not a measurement."""
    body: str | None = None
    unassessed_reason: str | None = None
    """Shown instead of body when body is None."""
    evidence: list[AssessmentEvidenceOut] = field(default_factory=list)
    claim: AnalysisClaimOut | None = None
    """The Sec 16 structured-reasoning claim mirroring this section's body
    text, if one exists (see _claim_for_text) - undefined for a NULL body
    or an assessment predating the claims layer, same as the web report's
    claimForText (AssessmentReport.tsx)."""
    group: str = "assessment"
    """Which of GROUP_INFO's three clusters this section prints under - the
    same context/assessment/notes grouping the web report uses (see
    REPORT_GROUPS in AssessmentReport.tsx). Both renderers watch for this
    changing as they walk build_report_sections()'s list and insert a new
    group heading whenever it does, rather than storing the groups as a
    separate nested structure - keeping one flat, ordered list is what lets
    build_report_sections() stay "the one place that decides what goes in
    the report" per this module's docstring."""


# Mirrors REPORT_GROUPS in AssessmentReport.tsx: same three clusters, same
# index/title/description, so a reader moving between the web report and an
# export sees the same structure. "context" isn't a ReportSection group (the
# export's "Input" block predates build_report_sections() and isn't in the
# list this dict keys off), so both renderers print it explicitly.
GROUP_INFO: dict[str, tuple[str, str, str]] = {
    "context": ("01", "context", "What was submitted, and what literature it's being read against."),
    "assessment": (
        "02",
        "assessment",
        "Grounded judgements - each one counted by how many real passages support it.",
    ),
    "notes": (
        "03",
        "notes",
        "What this reading doesn't settle, and how the recommendation was reached.",
    ),
}


def _claim_for_text(claims: list[AnalysisClaimOut], text: str | None) -> AnalysisClaimOut | None:
    """Same exact-match lookup as the web report's claimForText
    (AssessmentReport.tsx) - claim_text is written verbatim from the field
    (see assessment/claims.py), so exact match is reliable."""
    if not text:
        return None
    return next((c for c in claims if c.claim_text == text), None)


def build_report_sections(assessment: ResearchAssessmentOut) -> list[ReportSection]:
    by_role: dict[str, list[AssessmentEvidenceOut]] = {}
    for item in assessment.evidence:
        by_role.setdefault(item.role, []).append(item)

    applications_body = None
    applications_unassessed_reason = _APPLICATIONS_UNASSESSED_REASONS["not_assessed"]
    if assessment.potential_applications:
        applications_body = "\n".join(
            f"- {app['application']} (source: {app['source_paper']})" for app in assessment.potential_applications
        )
    elif assessment.potential_applications == []:
        applications_unassessed_reason = _APPLICATIONS_UNASSESSED_REASONS["no_evidence"]

    research_gap_body = assessment.research_gap_text
    if research_gap_body and assessment.research_gap_source:
        source_note = (
            "reused a reviewed candidate gap"
            if assessment.research_gap_source == "reused_candidate_gap"
            else "found for this input"
        )
        research_gap_body = f"{research_gap_body}\n({source_note})"
    research_gap_unassessed_reason = _GAP_UNASSESSED_REASONS.get(
        assessment.research_gap_source, _GAP_UNASSESSED_REASONS[None]
    )

    return [
        ReportSection(
            label="Existing solutions",
            body=assessment.comparison_summary,
            unassessed_reason="No retrieved paper had extracted claims to compare against.",
            evidence=by_role.get("comparison", []),
            claim=_claim_for_text(assessment.claims, assessment.comparison_summary),
        ),
        ReportSection(
            label="Novelty assessment",
            level=assessment.novelty_level,
            body=assessment.novelty_reasoning,
            unassessed_reason="Not enough evidence to judge novelty.",
            evidence=by_role.get("novelty", []),
            claim=_claim_for_text(assessment.claims, assessment.novelty_reasoning),
        ),
        ReportSection(
            label="Research gap",
            body=research_gap_body,
            unassessed_reason=research_gap_unassessed_reason,
            evidence=by_role.get("research_gap", []),
            claim=_claim_for_text(assessment.claims, assessment.research_gap_text),
        ),
        ReportSection(
            label="Potential applications",
            body=applications_body,
            unassessed_reason=applications_unassessed_reason,
            evidence=by_role.get("application", []),
        ),
        ReportSection(
            label="Product / technology opportunities",
            body=None,
            unassessed_reason=OPPORTUNITIES_REASON,
            evidence=by_role.get("opportunity", []),
        ),
        ReportSection(
            label="Technical feasibility",
            level=assessment.technical_feasibility_level,
            body=assessment.technical_feasibility_reasoning,
            unassessed_reason="Nothing close enough to ground a feasibility judgement.",
            evidence=by_role.get("feasibility", []),
            claim=_claim_for_text(assessment.claims, assessment.technical_feasibility_reasoning),
        ),
        ReportSection(
            label="Risks / limitations",
            body=assessment.risks_and_limitations,
            unassessed_reason="No retrieved paper stated a limitation.",
            evidence=by_role.get("risk", []),
            claim=_claim_for_text(assessment.claims, assessment.risks_and_limitations),
        ),
    ]


def _stats_tiles(assessment: ResearchAssessmentOut, related: list[RelatedPaper]) -> list[tuple[str, str]]:
    """The stats panel's four (label, value) tiles - confidence and
    human-reviewed are already on the assessment; evidence-quote and
    papers-cited counts are derived here rather than stored, since they're
    always recomputable from assessment.evidence."""
    return [
        ("confidence", assessment.confidence or "—"),
        ("evidence quotes", str(len(assessment.evidence))),
        ("papers cited", str(len(related))),
        ("human reviewed", "yes" if assessment.human_reviewed else "no"),
    ]


def _section_evidence_counts(sections: list[ReportSection]) -> list[tuple[str, int]]:
    return [(section.label, len(section.evidence)) for section in sections]


_LEVEL_RANK = {"low": 1, "medium": 2, "high": 3}


def _claim_suffix(claim: AnalysisClaimOut | None) -> str | None:
    if claim is None:
        return None
    return f"{claim.claim_type} · confidence: {claim.confidence}"


def _level_dots(level: str | None) -> str | None:
    """A filled/hollow-dot readout of a rule-based level (low/medium/high),
    for skimming the eyebrow line without reading the word. Bullet (U+2022)
    and middle dot (U+00B7) are both in WinAnsi/Latin-1, so they render in
    every font this module uses - no risk of a missing glyph."""
    rank = _LEVEL_RANK.get(level or "")
    if rank is None:
        return None
    return "•" * rank + "·" * (3 - rank)


@dataclass
class RelatedPaper:
    paper_id: str
    title: str
    link: str | None


def _source_link(url: str | None, doi: str | None) -> str | None:
    if url:
        return url
    if doi:
        return f"https://doi.org/{doi}"
    return None


def _related_papers(assessment: ResearchAssessmentOut) -> list[RelatedPaper]:
    """Every distinct paper cited as evidence, in first-seen order - this is
    the report's reference list (rendered at the end), and also where each
    evidence quote's citation looks up a link to hyperlink to."""
    seen: dict[str, RelatedPaper] = {}
    for item in assessment.evidence:
        pid = str(item.paper_id)
        if pid not in seen:
            seen[pid] = RelatedPaper(
                paper_id=pid, title=item.paper_title, link=_source_link(item.paper_url, item.paper_doi)
            )
    return list(seen.values())


def _docx_rgb(hex_color: str):
    from docx.shared import RGBColor

    return RGBColor.from_string(hex_color.lstrip("#"))


def _docx_bottom_border(paragraph, hex_color: str) -> None:
    """python-docx has no paragraph-border API, so this drops down to the
    underlying XML - the same recipe used across the python-docx ecosystem
    for a bare horizontal rule (there's no built-in flowable for one)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    borders.append(bottom)
    p_pr.append(borders)


def _docx_eyebrow(document, text: str, *, color: str = INK_FAINT, heading: bool = False, dots: str | None = None):
    """A small uppercase label, standing in for the web report's `.eyebrow`
    class - Space Grotesk is referenced by name only (Word substitutes if a
    reader doesn't have it installed), unlike the PDF path which embeds it.

    heading=True gives the paragraph Word's built-in "Heading 2" style on
    top of the same direct run formatting - direct formatting always wins
    visually, but the style is what makes the paragraph show up in Word's
    Navigation pane, which is the report's table of contents."""
    from docx.shared import Pt

    p = document.add_paragraph(style="Heading 2" if heading else None)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Space Grotesk"
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = _docx_rgb(color)
    if dots:
        dots_run = p.add_run("  " + dots)
        dots_run.font.name = "Source Serif 4"
        dots_run.font.size = Pt(9)
        dots_run.font.bold = False
        dots_run.font.color.rgb = _docx_rgb(INK_SOFT)
    return p


def _docx_group_heading(document, index: str, title: str, description: str) -> None:
    """One of the report's three cluster headings (context/assessment/notes)
    - bigger and bolder than _docx_eyebrow's per-field labels, and given
    Word's "Heading 1" style so it nests above them in the Navigation pane:
    Heading 1 group, Heading 2 fields beneath it, same hierarchy the web
    report's jump-nav shows."""
    from docx.shared import Pt

    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(20)
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(f"{index}  {title.upper()}")
    run.font.name = "Space Grotesk"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = _docx_rgb(INK)

    desc = document.add_paragraph()
    desc.paragraph_format.space_after = Pt(8)
    desc_run = desc.add_run(description)
    desc_run.font.name = "Source Serif 4"
    desc_run.font.size = Pt(9)
    desc_run.font.italic = True
    desc_run.font.color.rgb = _docx_rgb(INK_FAINT)


def _docx_hyperlink(paragraph, text: str, url: str, *, color: str = INK_SOFT, size_pt: float | None = None) -> None:
    """python-docx has no hyperlink API, so this drops to the underlying
    XML - the standard recipe for a `w:hyperlink` run pointing at an
    external relationship (there's no flowable for one, same situation as
    _docx_bottom_border above)."""
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color.lstrip("#"))
    run_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.append(underline)
    if size_pt is not None:
        size_el = OxmlElement("w:sz")
        size_el.set(qn("w:val"), str(int(size_pt * 2)))
        run_pr.append(size_el)
    run.append(run_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_stats_table(document, assessment: ResearchAssessmentOut, related: list[RelatedPaper]) -> None:
    """The confidence/evidence/papers/reviewed stats panel, as a 4-column,
    borderless table - each cell is a small label over a large value, same
    "stat tile" idea as the PDF path's _pdf_stats_table."""
    from docx.shared import Pt

    tiles = _stats_tiles(assessment, related)
    table = document.add_table(rows=2, cols=len(tiles))
    table.autofit = True
    for col, (label, value) in enumerate(tiles):
        label_run = table.rows[0].cells[col].paragraphs[0].add_run(label.upper())
        label_run.font.name = "Space Grotesk"
        label_run.font.size = Pt(7)
        label_run.font.bold = True
        label_run.font.color.rgb = _docx_rgb(INK_FAINT)

        value_run = table.rows[1].cells[col].paragraphs[0].add_run(value)
        value_run.font.name = "Space Grotesk"
        value_run.font.size = Pt(14)
        value_run.font.bold = True
        value_run.font.color.rgb = _docx_rgb(INK)


def _docx_comparison_summary(document, text: str) -> None:
    """Same fix as the PDF path's comparison_summary() closure: render each
    block/claim as its own paragraph instead of one add_paragraph(text) call
    - python-docx's add_t() stores a literal "\\n" as ordinary text (it does
    not emit a <w:br/>), so Word does not treat it as a line break either."""
    from docx.shared import Pt

    for block in text.split("\n\n"):
        if not block:
            continue
        heading, *claim_lines = block.split("\n")
        heading_p = document.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(8)
        heading_p.paragraph_format.space_after = Pt(2)
        heading_run = heading_p.add_run(heading)
        heading_run.font.name = "Space Grotesk"
        heading_run.font.bold = True
        heading_run.font.size = Pt(10.5)
        heading_run.font.color.rgb = _docx_rgb(INK)

        for line in claim_lines:
            match = _COMPARISON_CLAIM_RE.match(line)
            if not match:
                continue
            paper_title, claim_text = match.group(1), match.group(2)

            claim_p = document.add_paragraph()
            claim_p.paragraph_format.left_indent = Pt(14)
            claim_p.paragraph_format.space_after = Pt(0)
            claim_run = claim_p.add_run(claim_text)
            claim_run.font.color.rgb = _docx_rgb(INK_SOFT)

            source_p = document.add_paragraph()
            source_p.paragraph_format.left_indent = Pt(14)
            source_p.paragraph_format.space_after = Pt(8)
            source_run = source_p.add_run(paper_title.upper())
            source_run.font.name = "Space Grotesk"
            source_run.font.bold = True
            source_run.font.size = Pt(7.5)
            source_run.font.color.rgb = _docx_rgb(INK_FAINT)


def _docx_footer(document, assessment: ResearchAssessmentOut, generated_at: datetime) -> None:
    from docx.shared import Pt

    section = document.sections[0]
    section.footer.is_linked_to_previous = False
    footer_p = section.footer.paragraphs[0]
    run = footer_p.add_run(
        f"Assessment {str(assessment.id)[:8]}  ·  Exported {generated_at:%Y-%m-%d %H:%M UTC}"
    )
    run.font.name = "Source Serif 4"
    run.font.size = Pt(8)
    run.font.color.rgb = _docx_rgb(INK_FAINT)


def build_docx(assessment: ResearchAssessmentOut) -> bytes:
    import docx
    from docx.shared import Inches, Pt

    document = docx.Document()

    normal = document.styles["Normal"]
    normal.font.name = "Source Serif 4"
    normal.font.size = Pt(11)
    normal.font.color.rgb = _docx_rgb(INK)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Research Assessment")
    title_run.font.name = "Space Grotesk"
    title_run.font.size = Pt(15)
    title_run.font.bold = True
    title_run.font.color.rgb = _docx_rgb(INK_FAINT)

    _docx_eyebrow(document, "recommendation")
    headline = document.add_paragraph()
    headline_run = headline.add_run(assessment.recommendation or "Not assessed")
    headline_run.font.name = "Space Grotesk"
    headline_run.font.size = Pt(24)
    headline_run.font.bold = True
    headline_run.font.color.rgb = _docx_rgb(INK)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run(
        f"confidence: {assessment.confidence or '—'}   ·   "
        f"human reviewed: {'yes' if assessment.human_reviewed else 'no'}"
    )
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = _docx_rgb(INK_SOFT)
    _docx_bottom_border(meta, RULE)

    related = _related_papers(assessment)
    link_by_paper_id = {paper.paper_id: paper.link for paper in related}

    _docx_stats_table(document, assessment, related)

    sections = build_report_sections(assessment)
    counts = _section_evidence_counts(sections)
    chart = evidence_bar_chart_png(counts, width_pt=460)
    if chart:
        chart_png, _height_pt = chart
        _docx_eyebrow(document, "evidence by section")
        document.add_picture(io.BytesIO(chart_png), width=Inches(6.0))

    _docx_group_heading(document, *GROUP_INFO["context"])
    _docx_eyebrow(document, "input")
    document.add_paragraph(assessment.research_input.raw_text)
    input_type = document.add_paragraph()
    input_type_run = input_type.add_run(f"Type: {assessment.research_input.input_type}")
    input_type_run.font.size = Pt(9.5)
    input_type_run.font.color.rgb = _docx_rgb(INK_FAINT)

    last_group: str | None = None
    for section in sections:
        if section.group != last_group:
            _docx_group_heading(document, *GROUP_INFO[section.group])
            last_group = section.group

        label = section.label
        dots = _level_dots(section.level)
        if section.level:
            label += f"  ·  {section.level.replace('_', ' ')}"
        claim_suffix = _claim_suffix(section.claim)
        if claim_suffix:
            label += f"  ·  {claim_suffix}"
        _docx_eyebrow(document, label, heading=True, dots=dots)

        if section.level:
            gauge_png = level_gauge_png(section.level)
            document.add_picture(io.BytesIO(gauge_png), width=Inches(1.9))

        if section.body and section.label == "Existing solutions":
            _docx_comparison_summary(document, section.body)
        elif section.body:
            document.add_paragraph(section.body)
        elif section.unassessed_reason:
            reason = document.add_paragraph()
            reason_run = reason.add_run(section.unassessed_reason)
            reason_run.italic = True
            reason_run.font.color.rgb = _docx_rgb(INK_FAINT)

        for item in section.evidence:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            quote_run = p.add_run(f"“{item.text}”  — ")
            quote_run.italic = True
            quote_run.font.color.rgb = _docx_rgb(INK_SOFT)
            link = link_by_paper_id.get(str(item.paper_id))
            if link:
                _docx_hyperlink(p, item.paper_title, link, size_pt=9)
            else:
                source_run = p.add_run(item.paper_title)
                source_run.font.size = Pt(9)
                source_run.font.color.rgb = _docx_rgb(INK_FAINT)
            if item.section:
                section_run = p.add_run(f" ({item.section})")
                section_run.font.size = Pt(9)
                section_run.font.color.rgb = _docx_rgb(INK_FAINT)

    if related:
        _docx_eyebrow(document, "references", heading=True)
        for i, paper in enumerate(related, start=1):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            number_run = p.add_run(f"{i}. ")
            number_run.font.size = Pt(9.5)
            number_run.font.color.rgb = _docx_rgb(INK_FAINT)
            if paper.link:
                _docx_hyperlink(p, paper.title, paper.link, color=INK, size_pt=9.5)
            else:
                title_run = p.add_run(paper.title)
                title_run.font.size = Pt(9.5)
                title_run.font.color.rgb = _docx_rgb(INK)

    _docx_footer(document, assessment, datetime.now(timezone.utc))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_PDF_FONTS_REGISTERED = False


def _register_pdf_fonts() -> None:
    """Registers the bundled static font instances with reportlab. Idempotent
    and cheap to call per-export - reportlab has no "is this font already
    registered" check of its own, so a module-level flag avoids re-reading
    the font files off disk on every call."""
    global _PDF_FONTS_REGISTERED
    if _PDF_FONTS_REGISTERED:
        return

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdfmetrics.registerFont(TTFont("SpaceGrotesk-Bold", str(FONTS_DIR / "SpaceGrotesk-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("SourceSerif4", str(FONTS_DIR / "SourceSerif4-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("SourceSerif4-Italic", str(FONTS_DIR / "SourceSerif4-Italic.ttf")))
    _PDF_FONTS_REGISTERED = True


def _pdf_styles():
    from reportlab.lib.styles import ParagraphStyle

    return {
        "title": ParagraphStyle(
            "title", fontName="SpaceGrotesk-Bold", fontSize=13, textColor=INK_FAINT, spaceAfter=2
        ),
        "headline": ParagraphStyle(
            "headline", fontName="SpaceGrotesk-Bold", fontSize=26, textColor=INK, leading=30, spaceAfter=6
        ),
        "meta": ParagraphStyle("meta", fontName="SourceSerif4", fontSize=9.5, textColor=INK_SOFT),
        "eyebrow": ParagraphStyle(
            "eyebrow", fontName="SpaceGrotesk-Bold", fontSize=9, textColor=INK_FAINT, spaceBefore=16, spaceAfter=4
        ),
        # The report's three cluster headings (context/assessment/notes) -
        # bigger and bolder than "eyebrow"/"section", and given their own
        # style name so afterFlowable can bookmark them at outline level 0
        # with the per-field "section" headings nested at level 1 beneath.
        "group": ParagraphStyle("group", fontName="SpaceGrotesk-Bold", fontSize=13, textColor=INK, spaceBefore=22, spaceAfter=2),
        "group_description": ParagraphStyle(
            "group_description", fontName="SourceSerif4-Italic", fontSize=9, textColor=INK_FAINT, spaceAfter=10
        ),
        # Same look as "eyebrow" - a distinct style name so afterFlowable
        # (see _AssessmentDocTemplate) can tell a report section heading
        # apart from a meta label like "Input" and only bookmark the former.
        "section": ParagraphStyle(
            "section", fontName="SpaceGrotesk-Bold", fontSize=9, textColor=INK_FAINT, spaceBefore=16, spaceAfter=4
        ),
        "body": ParagraphStyle("body", fontName="SourceSerif4", fontSize=10.5, textColor=INK, leading=15),
        "reason": ParagraphStyle(
            "reason", fontName="SourceSerif4-Italic", fontSize=10, textColor=INK_FAINT, leading=14
        ),
        "quote": ParagraphStyle(
            "quote",
            fontName="SourceSerif4-Italic",
            fontSize=9.5,
            textColor=INK_SOFT,
            leading=13,
            leftIndent=12,
            spaceAfter=3,
        ),
        "quote_source": ParagraphStyle(
            "quote_source", fontName="SourceSerif4", fontSize=8, textColor=INK_FAINT, leftIndent=12, spaceAfter=8
        ),
        "bullet": ParagraphStyle(
            "bullet", fontName="SourceSerif4", fontSize=10.5, textColor=INK, leading=15, leftIndent=12, spaceAfter=2
        ),
        "comparison_heading": ParagraphStyle(
            "comparison_heading", fontName="SpaceGrotesk-Bold", fontSize=9.5, textColor=INK, leading=13, spaceAfter=4
        ),
        "comparison_claim": ParagraphStyle(
            "comparison_claim",
            fontName="SourceSerif4",
            fontSize=10,
            textColor=INK_SOFT,
            leading=14,
            leftIndent=12,
            spaceAfter=1,
        ),
        "comparison_source": ParagraphStyle(
            "comparison_source",
            fontName="SpaceGrotesk-Bold",
            fontSize=7.5,
            textColor=INK_FAINT,
            leftIndent=12,
            spaceAfter=8,
        ),
        # Two-line markup (label <br/> value) for each stats-panel tile -
        # the two sizes/colors come from inline <font> tags in the markup
        # itself, so this base style only needs to set the leading that
        # keeps those two lines readable.
        "stat_tile": ParagraphStyle("stat_tile", fontName="SourceSerif4", fontSize=7, leading=18),
    }


def _make_pdf_footer(assessment: ResearchAssessmentOut, generated_at: datetime):
    """A page-number-and-provenance footer. Built as a closure (rather than
    a plain onFirstPage/onLaterPages callback) because SimpleDocTemplate
    only ever passes it (canvas, doc) - the assessment id and export
    timestamp have to come from here instead."""

    def footer(canvas, doc) -> None:
        from reportlab.lib.colors import HexColor

        canvas.saveState()
        canvas.setStrokeColor(HexColor(RULE))
        canvas.setLineWidth(0.5)
        canvas.line(doc.leftMargin, 0.6 * 72, doc.pagesize[0] - doc.rightMargin, 0.6 * 72)
        canvas.setFont("SourceSerif4", 8)
        canvas.setFillColor(HexColor(INK_FAINT))
        canvas.drawString(
            doc.leftMargin,
            0.4 * 72,
            f"Assessment {str(assessment.id)[:8]}  ·  Exported {generated_at:%Y-%m-%d %H:%M UTC}",
        )
        canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 0.4 * 72, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    return footer


class _AssessmentDocTemplate:
    """Mixed in over SimpleDocTemplate (built lazily in build_pdf, since the
    class itself is only importable once reportlab is) to add a PDF outline
    (bookmarks) entry for each report section - the same "read structure I
    already lay out" idea as the docx Heading 2 style in _docx_eyebrow."""

    _bookmark_counter = 0

    def afterFlowable(self, flowable) -> None:
        style = getattr(flowable, "style", None)
        if style is None or style.name not in ("group", "section"):
            return
        level = 0 if style.name == "group" else 1
        key = f"section-{self._bookmark_counter}"
        self._bookmark_counter += 1
        self.canv.bookmarkPage(key)
        self.canv.addOutlineEntry(flowable.getPlainText(), key, level=level, closed=False)


def _pdf_stats_table(assessment: ResearchAssessmentOut, related: list[RelatedPaper], styles, content_width: float):
    """The confidence/evidence/papers/reviewed stats panel, as a borderless
    4-column Table - each cell is a two-line Paragraph (label over value),
    same "stat tile" idea as the docx path's _docx_stats_table."""
    from reportlab.platypus import Paragraph, Table, TableStyle

    tiles = _stats_tiles(assessment, related)
    cells = [
        Paragraph(
            f'<font name="SpaceGrotesk-Bold" size="7" color="{INK_FAINT}">{_escape(label.upper())}</font>'
            f'<br/><font name="SpaceGrotesk-Bold" size="14" color="{INK}">{_escape(value)}</font>',
            styles["stat_tile"],
        )
        for label, value in tiles
    ]
    table = Table([cells], colWidths=[content_width / len(cells)] * len(cells))
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    return table


def _pdf_link(text: str, url: str | None, *, color: str) -> str:
    """reportlab's `<link>` tag (not HTML's `<a>`) is the documented way to
    make part of a Paragraph clickable - it's the one markup tag guaranteed
    to accept a `color` attribute across reportlab versions."""
    escaped = _escape(text)
    if not url:
        return escaped
    return f'<link href="{_escape(url)}" color="{color}">{escaped}</link>'


def build_pdf(assessment: ResearchAssessmentOut) -> bytes:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import LETTER
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )
    from reportlab.platypus.flowables import HRFlowable

    class _DocTemplate(_AssessmentDocTemplate, SimpleDocTemplate):
        pass

    _register_pdf_fonts()
    styles = _pdf_styles()
    story = []
    margin = 0.9 * 72
    content_width = LETTER[0] - 2 * margin

    def para(text: str, style: str, *, markup: str | None = None) -> None:
        story.append(Paragraph(markup if markup is not None else _escape(text), styles[style]))

    def comparison_summary(text: str) -> None:
        """comparison_summary is pre-formatted into blocks (heading, then
        "- "paper": claim" lines - assessment/existing_solutions.py). A plain
        para(text, "body") call collapses all of that onto one flowed
        paragraph, since reportlab treats "\\n" as ordinary whitespace, not a
        line break - so each block/line needs its own Paragraph, the same
        structure the web report's ComparisonSummary component renders."""
        for block in text.split("\n\n"):
            if not block:
                continue
            heading, *claim_lines = block.split("\n")
            para(heading, "comparison_heading")
            for line in claim_lines:
                match = _COMPARISON_CLAIM_RE.match(line)
                if not match:
                    continue
                paper_title, claim_text = match.group(1), match.group(2)
                para(claim_text, "comparison_claim")
                para(paper_title.upper(), "comparison_source")

    def rule() -> None:
        story.append(HRFlowable(width="100%", thickness=0.75, color=HexColor(RULE), spaceBefore=6, spaceAfter=14))

    def group_heading(index: str, title: str, description: str) -> None:
        para(f"{index}  {title.upper()}", "group")
        para(description, "group_description")

    related = _related_papers(assessment)
    link_by_paper_id = {paper.paper_id: paper.link for paper in related}
    sections = build_report_sections(assessment)
    counts = _section_evidence_counts(sections)

    para("RESEARCH ASSESSMENT", "title")
    para(assessment.recommendation or "Not assessed", "headline")
    para(
        f"confidence: {assessment.confidence or '—'}"
        f"   ·   human reviewed: {'yes' if assessment.human_reviewed else 'no'}",
        "meta",
    )
    rule()

    story.append(_pdf_stats_table(assessment, related, styles, content_width))
    story.append(Spacer(1, 16))

    chart = evidence_bar_chart_png(counts, width_pt=content_width)
    if chart:
        chart_png, chart_height = chart
        para("Evidence by section", "eyebrow")
        story.append(Image(io.BytesIO(chart_png), width=content_width, height=chart_height))
        story.append(Spacer(1, 10))

    group_heading(*GROUP_INFO["context"])
    para("Input", "eyebrow")
    para(assessment.research_input.raw_text, "body")
    para(f"Type: {assessment.research_input.input_type}", "meta")

    if assessment.research_gap_text:
        para("Gap", "eyebrow")
        para(assessment.research_gap_text.splitlines()[0], "body")

    story.append(PageBreak())

    last_group: str | None = None
    for section in sections:
        if section.group != last_group:
            group_heading(*GROUP_INFO[section.group])
            last_group = section.group

        heading = _escape(section.label)
        if section.level:
            heading += f"  ·  {section.level.replace('_', ' ')}"
        claim_suffix = _claim_suffix(section.claim)
        if claim_suffix:
            heading += f"  ·  {_escape(claim_suffix)}"
        dots = _level_dots(section.level)
        if dots:
            heading += f'  <font name="SourceSerif4" size="9" color="{INK_SOFT}">{dots}</font>'
        para(section.label, "section", markup=heading)

        if section.level:
            story.append(Image(io.BytesIO(level_gauge_png(section.level)), width=140, height=10))
            story.append(Spacer(1, 6))

        if section.body and section.label == "Existing solutions":
            comparison_summary(section.body)
        elif section.body:
            para(section.body, "body")
        elif section.unassessed_reason:
            para(section.unassessed_reason, "reason")

        for item in section.evidence:
            suffix = f" ({_escape(item.section)})" if item.section else ""
            para(f"“{item.text}”", "quote")
            link = link_by_paper_id.get(str(item.paper_id))
            source_markup = f"— {_pdf_link(item.paper_title, link, color=INK_FAINT)}{suffix}"
            para(item.paper_title, "quote_source", markup=source_markup)

    if related:
        para("References", "section")
        for i, paper in enumerate(related, start=1):
            markup = f"{i}. {_pdf_link(paper.title, paper.link, color=INK)}"
            para(paper.title, "bullet", markup=markup)

    buffer = io.BytesIO()
    doc = _DocTemplate(buffer, pagesize=LETTER, topMargin=margin, bottomMargin=margin, leftMargin=margin, rightMargin=margin)
    generated_at = datetime.now(timezone.utc)
    footer = _make_pdf_footer(assessment, generated_at)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return buffer.getvalue()


def _escape(text: str) -> str:
    """reportlab's Paragraph interprets its text as a small XML/markup dialect."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_MD_INLINE_SPECIAL = re.compile(r"([\\`*_\[\]])")
_MD_LINE_START_MARKER = re.compile(r"(?m)^([-*+#>]|\d+\.)(?=\s|$)")


def _md_escape(text: str) -> str:
    """Escapes characters CommonMark would otherwise treat as markup, so a
    field's actual text (extracted from real paper prose, never authored
    as markdown) renders as plain text instead of accidentally toggling
    emphasis, a link, or - if it happens to start a line with "-", "#",
    "1.", etc. - a list/heading/blockquote it never meant to be. Only
    escapes what would actually trigger: backslash/backtick/asterisk/
    underscore/brackets can flip into markup anywhere inline, but a
    hyphen or period is only special as a line-starting list/heading
    marker - escaping every one mid-sentence (a real sentence full of
    hyphenated words and periods) would bury the text in backslashes for
    no reason, defeating the point of a format meant to stay readable
    unrendered."""
    escaped = _MD_INLINE_SPECIAL.sub(r"\\\1", text)
    return _MD_LINE_START_MARKER.sub(r"\\\1", escaped)


def _md_comparison_summary(text: str) -> list[str]:
    """Same block/claim-line structure as _docx_comparison_summary and the
    PDF path's comparison_summary() closure - comparison_summary is
    pre-formatted into blocks (heading, then '- "paper": claim' lines, see
    assessment/existing_solutions.py), not something a generic paragraph
    renderer can flow correctly."""
    lines: list[str] = []
    for block in text.split("\n\n"):
        if not block:
            continue
        heading, *claim_lines = block.split("\n")
        lines.append(f"**{_md_escape(heading)}**")
        lines.append("")
        for line in claim_lines:
            match = _COMPARISON_CLAIM_RE.match(line)
            if not match:
                continue
            paper_title, claim_text = match.group(1), match.group(2)
            lines.append(f"> {_md_escape(claim_text)}")
            lines.append(f"> — *{_md_escape(paper_title)}*")
            lines.append("")
    return lines


def build_markdown(assessment: ResearchAssessmentOut) -> bytes:
    """Plain-text-first export: no charts or gauges (nothing here has a
    text rendering worth the effort - a reader wants the numbers, not a
    redrawn bar), but everything build_report_sections() decides belongs
    in the report is present, in the same order, same as the docx/pdf
    paths. Meant for pasting into an issue tracker, a wiki page, or
    anywhere else prose beats a binary file - the one export format a
    reader can diff, grep, or read without opening Word or a PDF viewer.
    """
    related = _related_papers(assessment)
    link_by_paper_id = {paper.paper_id: paper.link for paper in related}
    sections = build_report_sections(assessment)

    lines: list[str] = ["# Research Assessment", ""]
    lines.append(f"## {_md_escape(assessment.recommendation or 'Not assessed')}")
    lines.append("")
    lines.append(
        f"confidence: {assessment.confidence or '—'} · "
        f"human reviewed: {'yes' if assessment.human_reviewed else 'no'}"
    )
    lines.append("")

    for label, value in _stats_tiles(assessment, related):
        lines.append(f"- **{label}**: {_md_escape(value)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    last_group: str | None = None

    def group_heading(group: str) -> None:
        index, title, description = GROUP_INFO[group]
        lines.append(f"## {index}  {title.upper()}")
        lines.append("")
        lines.append(f"*{_md_escape(description)}*")
        lines.append("")

    group_heading("context")
    lines.append("### Input")
    lines.append("")
    lines.append(_md_escape(assessment.research_input.raw_text))
    lines.append("")
    lines.append(f"Type: {assessment.research_input.input_type}")
    lines.append("")
    last_group = "context"

    for section in sections:
        if section.group != last_group:
            group_heading(section.group)
            last_group = section.group

        heading = f"### {_md_escape(section.label)}"
        if section.level:
            heading += f" · {section.level.replace('_', ' ')}"
        claim_suffix = _claim_suffix(section.claim)
        if claim_suffix:
            heading += f" · {_md_escape(claim_suffix)}"
        lines.append(heading)
        lines.append("")

        if section.body and section.label == "Existing solutions":
            lines.extend(_md_comparison_summary(section.body))
        elif section.body:
            lines.append(_md_escape(section.body))
            lines.append("")
        elif section.unassessed_reason:
            lines.append(f"*{_md_escape(section.unassessed_reason)}*")
            lines.append("")

        for item in section.evidence:
            lines.append(f'> "{_md_escape(item.text)}"')
            link = link_by_paper_id.get(str(item.paper_id))
            source = f"[{_md_escape(item.paper_title)}]({link})" if link else _md_escape(item.paper_title)
            suffix = f" ({_md_escape(item.section)})" if item.section else ""
            lines.append(f"> — {source}{suffix}")
            lines.append("")

    if related:
        lines.append("### References")
        lines.append("")
        for i, paper in enumerate(related, start=1):
            entry = f"[{_md_escape(paper.title)}]({paper.link})" if paper.link else _md_escape(paper.title)
            lines.append(f"{i}. {entry}")
        lines.append("")

    lines.append("---")
    generated_at = datetime.now(timezone.utc)
    lines.append(f"Assessment {str(assessment.id)[:8]} · Exported {generated_at:%Y-%m-%d %H:%M} UTC")
    lines.append("")

    return "\n".join(lines).encode("utf-8")
